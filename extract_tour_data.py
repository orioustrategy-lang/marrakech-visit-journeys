import os
import json
from pathlib import Path

# Try to import required libraries
try:
    import PyPDF2
    import docx
except ImportError:
    print("Installing required packages...")
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "PyPDF2", "python-docx"])
    import PyPDF2
    import docx

def extract_pdf_text(pdf_path):
    """Extract text from a PDF file."""
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        print(f"Error reading PDF {pdf_path}: {e}")
        return ""

def extract_docx_text(docx_path):
    """Extract text from a DOCX file."""
    try:
        doc = docx.Document(docx_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        
        return text
    except Exception as e:
        print(f"Error reading DOCX {docx_path}: {e}")
        return ""

def parse_tour_info(text, filename, category):
    """Parse tour information from extracted text."""
    lines = text.split('\n')
    
    # Initialize data structure
    tour_data = {
        "filename": filename,
        "category": category,
        "title": "",
        "description": "",
        "duration": "",
        "group_size": "",
        "price": "",
        "highlights": [],
        "whats_included": [],
        "itinerary": [],
        "raw_text": text  # Include raw text for manual review
    }
    
    # Try to extract title (usually first non-empty line or after "title:")
    for line in lines[:10]:
        if line.strip() and not line.strip().startswith(('http', 'www')):
            tour_data["title"] = line.strip()
            break
    
    # Search for common patterns
    text_lower = text.lower()
    
    # Duration
    if 'duration:' in text_lower:
        idx = text_lower.find('duration:')
        duration_text = text[idx:idx+200].split('\n')[0]
        tour_data["duration"] = duration_text.replace('Duration:', '').replace('duration:', '').strip()
    elif 'day' in text_lower or 'hour' in text_lower:
        for line in lines:
            if 'day' in line.lower() or 'hour' in line.lower():
                if len(line) < 100:  # Likely a duration line
                    tour_data["duration"] = line.strip()
                    break
    
    # Group size
    for keyword in ['group size:', 'participants:', 'maximum:', 'max group']:
        if keyword in text_lower:
            idx = text_lower.find(keyword)
            size_text = text[idx:idx+200].split('\n')[0]
            tour_data["group_size"] = size_text.split(':')[-1].strip()
            break
    
    # Price
    for keyword in ['price:', 'cost:', 'from €', 'from $', '€', 'price per']:
        if keyword in text_lower:
            for line in lines:
                if any(symbol in line for symbol in ['€', '$', 'EUR', 'USD']) and len(line) < 150:
                    tour_data["price"] = line.strip()
                    break
            break
    
    # Highlights
    if 'highlights' in text_lower or 'what to expect' in text_lower:
        in_highlights = False
        for line in lines:
            line_lower = line.lower()
            if 'highlight' in line_lower or 'what to expect' in line_lower:
                in_highlights = True
                continue
            if in_highlights:
                if line.strip() and (line.strip().startswith(('•', '-', '*', '▸')) or line.strip()[0].isdigit()):
                    tour_data["highlights"].append(line.strip())
                elif line.strip() and len(line.strip()) > 20:
                    if 'include' in line_lower or 'itinerary' in line_lower:
                        break
                    tour_data["highlights"].append(line.strip())
                elif len(tour_data["highlights"]) > 3 and not line.strip():
                    break
    
    # What's included
    if "what's included" in text_lower or "included" in text_lower or "includes:" in text_lower:
        in_included = False
        for line in lines:
            line_lower = line.lower()
            if "what's included" in line_lower or "includes:" in line_lower:
                in_included = True
                continue
            if in_included:
                if line.strip() and (line.strip().startswith(('•', '-', '*', '▸', '✓')) or line.strip()[0].isdigit()):
                    tour_data["whats_included"].append(line.strip())
                elif 'not included' in line_lower or 'itinerary' in line_lower:
                    break
                elif len(tour_data["whats_included"]) > 5 and not line.strip():
                    break
    
    # Itinerary
    if 'itinerary' in text_lower or 'program' in text_lower or 'schedule' in text_lower:
        in_itinerary = False
        for line in lines:
            line_lower = line.lower()
            if 'itinerary' in line_lower or 'program' in line_lower or 'schedule' in line_lower:
                in_itinerary = True
                continue
            if in_itinerary:
                if line.strip():
                    tour_data["itinerary"].append(line.strip())
                elif len(tour_data["itinerary"]) > 5 and not line.strip():
                    break
    
    # Description - try to get the main description paragraph
    for i, line in enumerate(lines):
        if len(line.strip()) > 100 and not any(kw in line.lower() for kw in ['http', 'www', 'highlights', 'included']):
            tour_data["description"] = line.strip()
            # Get next few lines if they're part of the description
            for j in range(i+1, min(i+5, len(lines))):
                if len(lines[j].strip()) > 50:
                    tour_data["description"] += " " + lines[j].strip()
                else:
                    break
            break
    
    return tour_data

def process_directory(directory_path, category_name):
    """Process all PDF and DOCX files in a directory."""
    results = []
    directory = Path(directory_path)
    
    if not directory.exists():
        print(f"Directory not found: {directory_path}")
        return results
    
    # Get all PDF and DOCX files
    pdf_files = list(directory.glob('*.pdf'))
    docx_files = list(directory.glob('*.docx'))
    
    print(f"\nProcessing {category_name}:")
    print(f"  Found {len(pdf_files)} PDF files and {len(docx_files)} DOCX files")
    
    # Process PDF files
    for pdf_file in pdf_files:
        print(f"  - Processing: {pdf_file.name}")
        text = extract_pdf_text(str(pdf_file))
        tour_info = parse_tour_info(text, pdf_file.name, category_name)
        results.append(tour_info)
    
    # Process DOCX files
    for docx_file in docx_files:
        print(f"  - Processing: {docx_file.name}")
        text = extract_docx_text(str(docx_file))
        tour_info = parse_tour_info(text, docx_file.name, category_name)
        results.append(tour_info)
    
    return results

def main():
    # Define all directories to process
    base_path = r"C:\Users\pc\Desktop\freelance\marrakech-visit-journeys\src\assets\Marrakech Discover"
    
    directories = {
        "Activities": os.path.join(base_path, "Activities"),
        "Airport Transfer": os.path.join(base_path, "Airport Transfer"),
        "Day Trips": os.path.join(base_path, "Day Trips"),
        "Tours": os.path.join(base_path, "Tours"),
        "Trekking": os.path.join(base_path, "Treking")
    }
    
    all_tours = []
    
    # Process each directory
    for category, directory_path in directories.items():
        tours = process_directory(directory_path, category)
        all_tours.extend(tours)
    
    # Save to JSON file
    output_file = r"C:\Users\pc\Desktop\freelance\marrakech-visit-journeys\extracted_tour_data.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(all_tours, f, indent=2, ensure_ascii=False)
    
    print(f"\n✓ Extraction complete!")
    print(f"✓ Total tours extracted: {len(all_tours)}")
    print(f"✓ Data saved to: {output_file}")
    
    # Print summary
    print("\nSummary by category:")
    category_counts = {}
    for tour in all_tours:
        cat = tour['category']
        category_counts[cat] = category_counts.get(cat, 0) + 1
    
    for category, count in sorted(category_counts.items()):
        print(f"  {category}: {count} tours")

if __name__ == "__main__":
    main()
